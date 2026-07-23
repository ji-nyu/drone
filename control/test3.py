from __future__ import annotations

import argparse
import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


# ============================================================
# 1. 사용자 설정
# ============================================================

MODEL_PATH = Path(
    os.getenv(
        "MODEL_PATH",
        "D:\\model\\EXAONE-3.5-7.8B-Q4",
    )
)

OUTPUT_DIR = Path(__file__).resolve().parent / "report_test_results"
SURVEY_DIR = Path(__file__).resolve().parent / "logs" / "surveys"
MAX_NEW_TOKENS = 1600
MAX_CPU_MEMORY = "48GiB"

# 기관 정보: 실제 배포 시 반드시 수정
ORGANIZATION_NAME = os.getenv("REPORT_ORGANIZATION", "기관명")
DEPARTMENT_NAME = os.getenv("REPORT_DEPARTMENT", "담당 부서")
AUTHOR_NAME = os.getenv("REPORT_AUTHOR", "작성자")
REVIEWER_NAME = os.getenv("REPORT_REVIEWER", "검토자")
APPROVER_NAME = os.getenv("REPORT_APPROVER", "승인자")
DOCUMENT_SECURITY = os.getenv("REPORT_SECURITY", "내부 검토용")

# Windows Word에서는 맑은 고딕으로 자동 대체할 수 있다.
KOREAN_FONT = os.getenv("REPORT_FONT", "Noto Sans CJK KR")
TITLE_COLOR = "17365D"
SUBTITLE_COLOR = "365F91"
LIGHT_BLUE = "DCE6F1"
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9E2F3"
DARK_GRAY = "595959"
WHITE = "FFFFFF"


# ============================================================
# 2. LLM 출력 스키마 및 프롬프트
# ============================================================

SYSTEM_PROMPT = """
당신은 공공기관 및 연구기관의 해양환경 조사 보고서를 작성하는 전문 분석가다.
사용자가 제공한 JSON만 근거로 보고서 본문에 들어갈 분석 문장을 작성한다.

[절대 규칙]
1. 입력 JSON에 없는 날짜, 위치, 수량, 원인, 인력, 비용, 법적 판단을 생성하지 않는다.
2. 모든 수량과 고유명사를 원문 그대로 유지한다.
3. 확인되지 않은 원인은 단정하지 않는다.
4. 탐지 신뢰도는 객체 탐지 모델의 신뢰도일 뿐, 현장 확정 판정으로 표현하지 않는다.
5. 문체는 공식적이고 객관적인 행정 보고서 문체로 작성한다.
6. 마크다운, 표, 코드블록을 출력하지 않는다.
7. 아래 JSON 객체 하나만 출력한다. JSON 앞뒤에 설명을 붙이지 않는다.
8. collection_priority, recommended_actions, limitations 가 입력에 비어 있으면
   탐지 수량·구역 분포·신뢰도만으로 새로 작성한다.
9. 위 항목이 입력에 채워져 있으면 그 내용을 유지·다듬어 사용한다.

[출력 JSON 스키마]
{
  "executive_summary": "조사 목적과 핵심 탐지 결과를 요약한 3~5문장",
  "overall_assessment": "전체 탐지 결과의 분포와 특징을 분석한 2~4문장",
  "zone_analysis": [
    {
      "zone": "입력 데이터의 구역명",
      "analysis": "해당 구역의 탐지 결과를 객관적으로 분석한 1~3문장"
    }
  ],
  "collection_priority": "높음|중간|일반|현장 확인 필요 중 하나 (탐지 수량·분포 근거)",
  "priority_basis": "수거 우선순위의 근거를 설명한 2~3문장",
  "recommended_actions": ["탐지 수치에 근거한 조치 문장 2~4개"],
  "limitations": ["탐지 데이터·촬영 조건에 근거한 한계 문장 2~4개"],
  "conclusion": "종합 의견과 후속 확인 필요성을 정리한 3~5문장"
}
""".strip()


# ============================================================
# 3. 테스트 데이터
# ============================================================

TEST_CASES: dict[str, dict[str, Any]] = {
    "normal": {
        "mission_id": "MISSION-20260721-001",
        "inspection_date": "2026-07-21",
        "inspection_area": "제주시 해안 A구역",
        "flight_id": "TT-01-0721",
        "survey_method": "드론 영상 기반 객체 탐지",
        "survey_duration_minutes": 18,
        "total_detected_count": 18,
        "detections": [
            {
                "trash_type": "플라스틱병",
                "count": 9,
                "zone": "A-01",
                "average_confidence": 0.92,
            },
            {
                "trash_type": "비닐봉지",
                "count": 5,
                "zone": "A-02",
                "average_confidence": 0.88,
            },
            {
                "trash_type": "캔",
                "count": 4,
                "zone": "A-03",
                "average_confidence": 0.90,
            },
        ],
        "highest_density_zone": "A-01",
        "collection_priority": "높음",
        "recommended_actions": [
            "A-01 구역을 우선적으로 현장 확인한다.",
            "플라스틱병과 비닐봉지를 우선 수거한다.",
            "수거 완료 후 동일 구역을 재촬영한다.",
        ],
        "limitations": [
            "객체 추적을 통해 동일 객체의 중복 탐지를 제거했다.",
            "탐지 결과는 현장 확인을 통해 최종 검증해야 한다.",
        ],
    },
    "missing_information": {
        "mission_id": "MISSION-20260721-002",
        "inspection_date": "2026-07-21",
        "inspection_area": "제주시 해안 B구역",
        "flight_id": "TT-02-0721",
        "survey_method": "드론 영상 기반 객체 탐지",
        "survey_duration_minutes": None,
        "total_detected_count": 2,
        "detections": [
            {
                "trash_type": "스티로폼",
                "count": 2,
                "zone": None,
                "average_confidence": 0.76,
            }
        ],
        "highest_density_zone": None,
        "collection_priority": "현장 확인 필요",
        "recommended_actions": [],
        "limitations": [
            "세부 위치 정보가 기록되지 않았다.",
            "탐지 신뢰도가 상대적으로 낮아 현장 검증이 필요하다.",
        ],
    },
    "no_detection": {
        "mission_id": "MISSION-20260721-003",
        "inspection_date": "2026-07-21",
        "inspection_area": "제주시 해안 C구역",
        "flight_id": "TT-01-0721-02",
        "survey_method": "드론 영상 기반 객체 탐지",
        "survey_duration_minutes": 12,
        "total_detected_count": 0,
        "detections": [],
        "highest_density_zone": None,
        "collection_priority": "일반",
        "recommended_actions": [
            "현재 탐지 결과만으로 수거 지점을 지정하지 않는다.",
            "필요한 경우 다른 시간대에 추가 촬영한다.",
        ],
        "limitations": [
            "미탐지는 실제 쓰레기가 전혀 없음을 보장하지 않는다.",
            "촬영 각도와 조명 상태에 따라 탐지 결과가 달라질 수 있다.",
        ],
    },
}


# ============================================================
# 4. 공통 유틸리티
# ============================================================


def display_value(value: Any, suffix: str = "") -> str:
    if value is None or value == "":
        return "확인되지 않음"
    return f"{value}{suffix}"


def confidence_percent(value: Any) -> str:
    if value is None:
        return "확인되지 않음"
    try:
        return f"{float(value) * 100:.1f}%"
    except (TypeError, ValueError):
        return str(value)


def calculate_ratio(count: int, total: int) -> str:
    if total <= 0:
        return "0.0%"
    return f"{count / total * 100:.1f}%"


def remove_thinking_text(text: str) -> str:
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()


def extract_json_object(text: str) -> dict[str, Any]:
    cleaned = remove_thinking_text(text)
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned.strip(), flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned.strip())

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("모델 출력에서 JSON 객체를 찾지 못했습니다.")

    return json.loads(cleaned[start : end + 1])


def deterministic_narrative(data: dict[str, Any]) -> dict[str, Any]:
    """모델 출력 파싱 실패 시에도 보고서를 만들 수 있는 안전한 대체 본문."""

    total = int(data.get("total_detected_count") or 0)
    area = display_value(data.get("inspection_area"))
    date = display_value(data.get("inspection_date"))
    detections = data.get("detections", [])
    highest = data.get("highest_density_zone")

    if total > 0 and detections:
        type_text = ", ".join(
            f"{item.get('trash_type', '미분류')} {item.get('count', 0)}개"
            for item in detections
        )
        executive_summary = (
            f"{date} {area}에서 드론 영상 기반 해양 쓰레기 탐지를 수행하였다. "
            f"탐지 결과 총 {total}개의 객체가 확인되었으며, 종류별 현황은 {type_text}로 집계되었다. "
            "본 결과는 객체 탐지 모델의 분석 결과로서 현장 확인을 통한 최종 검증이 필요하다."
        )
        overall_assessment = (
            f"전체 탐지 객체는 총 {total}개이며, 구역별·종류별 분포에 차이가 확인되었다. "
            f"가장 높은 밀도 구역은 {display_value(highest)}로 기록되었다."
        )
    else:
        executive_summary = (
            f"{date} {area}에서 드론 영상 기반 해양 쓰레기 탐지를 수행하였다. "
            "해당 분석에서는 해양 쓰레기 객체가 탐지되지 않았다. "
            "다만 미탐지는 실제 현장에 쓰레기가 존재하지 않음을 보장하지 않으므로 촬영 조건과 탐지 범위를 함께 검토해야 한다."
        )
        overall_assessment = (
            "현재 탐지 결과만으로는 오염 지점을 특정하기 어렵다. "
            "필요한 경우 다른 시간대 또는 촬영 조건에서 추가 조사를 실시할 수 있다."
        )

    zone_analysis: list[dict[str, str]] = []
    for item in detections:
        zone = display_value(item.get("zone"))
        trash_type = display_value(item.get("trash_type"))
        count = int(item.get("count") or 0)
        confidence = confidence_percent(item.get("average_confidence"))
        zone_analysis.append(
            {
                "zone": zone,
                "analysis": (
                    f"{zone}에서 {trash_type} {count}개가 탐지되었으며, "
                    f"평균 탐지 신뢰도는 {confidence}로 기록되었다. "
                    "해당 결과는 현장 확인을 통해 최종 검증할 필요가 있다."
                ),
            }
        )

    if not zone_analysis:
        zone_analysis.append(
            {
                "zone": "전체 조사 구역",
                "analysis": "해당 조사에서는 구역별 탐지 객체가 기록되지 않았다.",
            }
        )

    # 입력에 비어 있으면 탐지 수치로 판단 항목 생성 (LLM 대체용)
    priority = data.get("collection_priority")
    if not priority:
        if total <= 0:
            priority = "일반"
        elif total >= 10:
            priority = "높음"
        elif total >= 3:
            priority = "중간"
        else:
            priority = "현장 확인 필요"

    actions = [str(a).strip() for a in (data.get("recommended_actions") or []) if str(a).strip()]
    if not actions:
        if total <= 0:
            actions = [
                "현재 탐지 결과만으로 수거 지점을 지정하지 않는다.",
                "필요한 경우 다른 시간대에 추가 촬영한다.",
            ]
        else:
            if highest:
                actions.append(f"{highest} 구역을 우선적으로 현장 확인한다.")
            top_types = sorted(
                detections, key=lambda d: int(d.get("count") or 0), reverse=True
            )[:2]
            if top_types:
                names = "·".join(str(d.get("trash_type") or "미분류") for d in top_types)
                actions.append(f"{names}을(를) 우선 수거 대상으로 검토한다.")
            actions.append("수거 또는 조치 전 현장 확인을 실시하고 동일 구역을 재촬영한다.")

    limitations = [str(a).strip() for a in (data.get("limitations") or []) if str(a).strip()]
    if not limitations:
        if total <= 0:
            limitations = [
                "미탐지는 실제 쓰레기가 전혀 없음을 보장하지 않는다.",
                "촬영 각도와 조명 상태에 따라 탐지 결과가 달라질 수 있다.",
            ]
        else:
            limitations = [
                "탐지 수량은 객체 추적 기반 집계이며 현장 실물 수와 다를 수 있다.",
                "탐지 신뢰도는 모델 점수일 뿐 현장 확정 판정이 아니다.",
                "조치 전 현장 확인이 필요하다.",
            ]

    return {
        "executive_summary": executive_summary,
        "overall_assessment": overall_assessment,
        "zone_analysis": zone_analysis,
        "collection_priority": priority,
        "priority_basis": (
            f"탐지 총 {total}개"
            + (f", 최고 밀도 구역 {highest}" if highest else "")
            + f"를 근거로 수거 우선순위를 '{priority}'로 판단하였다. "
            "실제 조치 여부와 순서는 현장 접근성, 안전성 및 현장 확인 결과를 종합하여 결정해야 한다."
        ),
        "recommended_actions": actions,
        "limitations": limitations,
        "conclusion": (
            "본 보고서는 드론 영상과 객체 탐지 결과를 기반으로 작성되었다. "
            "탐지 수량과 위치 정보는 조사 당시 입력 데이터를 기준으로 하며, 현장 상황과 차이가 있을 수 있다. "
            "따라서 수거 또는 행정 조치 전 현장 확인을 실시하고, 조치 완료 후 동일 구역을 재점검하는 것이 필요하다."
        ),
    }


def normalize_narrative(
    narrative: dict[str, Any],
    data: dict[str, Any],
) -> dict[str, Any]:
    fallback = deterministic_narrative(data)

    normalized = {
        "executive_summary": str(
            narrative.get("executive_summary") or fallback["executive_summary"]
        ).strip(),
        "overall_assessment": str(
            narrative.get("overall_assessment") or fallback["overall_assessment"]
        ).strip(),
        "collection_priority": str(
            narrative.get("collection_priority")
            or data.get("collection_priority")
            or fallback["collection_priority"]
        ).strip(),
        "priority_basis": str(
            narrative.get("priority_basis") or fallback["priority_basis"]
        ).strip(),
        "conclusion": str(
            narrative.get("conclusion") or fallback["conclusion"]
        ).strip(),
    }

    zone_analysis = narrative.get("zone_analysis")
    if not isinstance(zone_analysis, list) or not zone_analysis:
        zone_analysis = fallback["zone_analysis"]

    normalized["zone_analysis"] = [
        {
            "zone": str(item.get("zone") or "확인되지 않음"),
            "analysis": str(item.get("analysis") or "분석 내용 없음"),
        }
        for item in zone_analysis
        if isinstance(item, dict)
    ]

    for key in ("recommended_actions", "limitations"):
        value = narrative.get(key)
        if not isinstance(value, list) or not value:
            value = fallback[key]
        normalized[key] = [str(item).strip() for item in value if str(item).strip()]

    return normalized


# ============================================================
# 5. 실행 환경 및 GGUF 모델 로드
# ============================================================


def resolve_gguf_model_path(model_path: Path) -> Path:
    """폴더 또는 파일 경로를 실제 GGUF 파일 경로로 변환한다."""

    model_path = model_path.expanduser()

    if model_path.is_file():
        if model_path.suffix.lower() != ".gguf":
            raise ValueError(
                f"GGUF 모델 파일이 아닙니다: {model_path}\n"
                ".gguf 확장자의 모델 파일을 지정하세요."
            )
        return model_path.resolve()

    if not model_path.exists():
        raise FileNotFoundError(
            f"모델 경로를 찾을 수 없습니다: {model_path}\n"
            "--model-path 옵션 또는 MODEL_PATH 환경변수를 수정하세요."
        )

    if not model_path.is_dir():
        raise ValueError(f"지원하지 않는 모델 경로입니다: {model_path}")

    candidates = sorted(
        path for path in model_path.rglob("*.gguf")
        if path.is_file()
    )

    if not candidates:
        raise FileNotFoundError(
            f"폴더 안에서 GGUF 파일을 찾지 못했습니다: {model_path}\n"
            "Hugging Face에서 Q4_K_M 또는 IQ4_XS GGUF 파일을 다운로드하세요."
        )

    priority_keywords = (
        "Q4_K_M",
        "IQ4_XS",
        "Q5_K_M",
        "Q6_K",
        "Q8_0",
        "Q4_0",
    )

    for keyword in priority_keywords:
        matching = [
            path for path in candidates
            if keyword.lower() in path.name.lower()
        ]
        if matching:
            selected = matching[0]
            print(f"[안내] GGUF 파일 자동 선택: {selected.name}")
            return selected.resolve()

    if len(candidates) > 1:
        print("[경고] 여러 GGUF 파일 중 첫 번째 파일을 선택합니다.")
        for candidate in candidates:
            print(f"  - {candidate.name}")

    selected = candidates[0]
    print(f"[안내] GGUF 파일 자동 선택: {selected.name}")
    return selected.resolve()


def check_environment(
    model_path: Path,
    n_gpu_layers: int,
    n_ctx: int,
    n_batch: int,
) -> None:
    """GGUF 파일과 llama-cpp-python 실행 환경을 확인한다."""

    try:
        import llama_cpp
    except ImportError as error:
        raise RuntimeError(
            "llama-cpp-python이 설치되지 않았습니다.\n"
            "CPU 버전 설치: python -m pip install llama-cpp-python\n"
            "CUDA 버전은 공식 llama-cpp-python 안내에 따라 "
            "GGML_CUDA=on으로 설치하세요."
        ) from error

    file_size_gib = model_path.stat().st_size / (1024**3)

    print("=" * 72)
    print("실행 환경")
    print("=" * 72)
    print(f"GGUF 모델       : {model_path}")
    print(f"모델 파일 용량  : {file_size_gib:.2f} GiB")
    print(f"llama-cpp-python: {getattr(llama_cpp, '__version__', '확인 불가')}")
    print(f"GPU 오프로딩    : {n_gpu_layers}개 레이어")
    print(f"컨텍스트 길이   : {n_ctx}")
    print(f"배치 크기       : {n_batch}")
    print("=" * 72)


def load_model(
    model_path: Path,
    *,
    n_gpu_layers: int,
    n_ctx: int,
    n_batch: int,
    n_threads: int,
    chat_format: str | None,
    verbose: bool,
):
    """GGUF 모델을 llama.cpp 런타임으로 불러온다."""

    from llama_cpp import Llama

    resolved_threads = n_threads
    if resolved_threads <= 0:
        resolved_threads = max((os.cpu_count() or 4) - 2, 1)

    print("\nGGUF 모델을 불러오는 중...")
    print(f"모델 경로       : {model_path}")
    print(f"CPU 스레드      : {resolved_threads}")

    kwargs: dict[str, Any] = {
        "model_path": str(model_path),
        "n_ctx": n_ctx,
        "n_batch": n_batch,
        "n_gpu_layers": n_gpu_layers,
        "n_threads": resolved_threads,
        "seed": 42,
        "use_mmap": True,
        "use_mlock": False,
        "verbose": verbose,
    }

    if chat_format:
        kwargs["chat_format"] = chat_format

    try:
        model = Llama(**kwargs)
    except Exception as error:
        raise RuntimeError(
            "GGUF 모델 로드에 실패했습니다.\n"
            "CUDA 메모리 부족이면 --n-gpu-layers 값을 20에서 "
            "16, 12, 8 순으로 낮춰보세요.\n"
            f"원본 오류: {error}"
        ) from error

    print("GGUF 모델 로드 완료")
    return model


# ============================================================
# 6. 보고서 본문 생성
# ============================================================


def create_user_prompt(data: dict[str, Any]) -> str:
    return (
        "다음 해양 쓰레기 탐지 데이터만 근거로 기관 제출용 보고서 분석 문장을 작성하라.\n"
        "수량과 위치는 절대 변경하지 말고, 지정된 JSON 스키마만 출력하라.\n"
        "collection_priority, recommended_actions, limitations 가 비어 있으면 "
        "탐지 수량·구역·신뢰도를 근거로 직접 작성하라.\n\n"
        + json.dumps(data, ensure_ascii=False, indent=2)
    )


def generate_narrative(
    model,
    data: dict[str, Any],
) -> tuple[dict[str, Any], dict[str, Any]]:
    """GGUF 모델로 기관 보고서용 분석 문장을 생성한다."""

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": create_user_prompt(data)},
    ]

    request_options: dict[str, Any] = {
        "messages": messages,
        "max_tokens": MAX_NEW_TOKENS,
        "temperature": 0.0,
        "top_p": 1.0,
        "repeat_penalty": 1.04,
        "seed": 42,
    }

    started_at = time.perf_counter()
    json_mode_used = True

    try:
        response = model.create_chat_completion(
            **request_options,
            response_format={"type": "json_object"},
        )
    except Exception as error:
        print(
            "[경고] JSON 강제 모드 실패. 일반 채팅 생성으로 재시도합니다: "
            f"{error}"
        )
        json_mode_used = False
        response = model.create_chat_completion(**request_options)

    elapsed = time.perf_counter() - started_at

    choices = response.get("choices", [])
    if not choices:
        raise RuntimeError("모델 응답에 choices 항목이 없습니다.")

    choice = choices[0]
    message = choice.get("message") or {}
    raw_text = str(
        message.get("content")
        or choice.get("text")
        or ""
    ).strip()

    if not raw_text:
        raise RuntimeError("모델이 빈 응답을 생성했습니다.")

    try:
        parsed = extract_json_object(raw_text)
        parse_status = "success"
    except (ValueError, json.JSONDecodeError) as error:
        print(
            "[경고] 모델 JSON 파싱 실패. "
            f"안전한 기본 본문을 사용합니다: {error}"
        )
        parsed = deterministic_narrative(data)
        parse_status = "fallback"

    narrative = normalize_narrative(parsed, data)

    usage = response.get("usage", {}) or {}
    prompt_tokens = int(usage.get("prompt_tokens", 0) or 0)
    completion_tokens = int(usage.get("completion_tokens", 0) or 0)

    tokens_per_second = (
        completion_tokens / elapsed
        if elapsed > 0 and completion_tokens > 0
        else 0.0
    )

    metadata = {
        "elapsed_seconds": round(elapsed, 2),
        "input_tokens": prompt_tokens,
        "generated_tokens": completion_tokens,
        "tokens_per_second": round(tokens_per_second, 2),
        "parse_status": parse_status,
        "json_mode_used": json_mode_used,
        "finish_reason": choice.get("finish_reason"),
        "raw_model_output": raw_text,
        "runtime": "llama-cpp-python",
    }

    return narrative, metadata


# ============================================================
# 7. DOCX 서식 유틸리티
# ============================================================


def set_run_font(run, size: float = 10.5, bold: bool = False, color: str | None = None):
    run.font.name = KOREAN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), KOREAN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), KOREAN_FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), KOREAN_FONT)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_format(
    paragraph,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.45,
):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, margin_value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "A6A6A6", size: int = 6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 9.5,
    color: str | None = None,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_format(paragraph, alignment=alignment, after=0, line_spacing=1.1)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_column_widths(table, widths_cm: list[float]):
    for row in table.rows:
        for idx, width_cm in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width_cm)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_run_font(run, size=8.5, color=DARK_GRAY)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    field_run = paragraph.add_run()
    field_run._r.append(begin)
    field_run._r.append(instr)
    field_run._r.append(separate)
    field_run._r.append(text)
    field_run._r.append(end)


def add_section_heading(document: Document, number: str, title: str):
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, before=8, after=5, line_spacing=1.0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(f"{number}. {title}")
    set_run_font(run, size=14, bold=True, color=TITLE_COLOR)

    p_pr = paragraph._p.get_or_add_pPr()
    bottom = OxmlElement("w:pBdr")
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), "10")
    border.set(qn("w:space"), "2")
    border.set(qn("w:color"), TITLE_COLOR)
    bottom.append(border)
    p_pr.append(bottom)


def add_body_paragraph(document: Document, text: str, indent: bool = False):
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.7)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    return paragraph


def add_numbered_items(document: Document, items: list[str]):
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=4)
        paragraph.paragraph_format.left_indent = Cm(0.4)
        paragraph.paragraph_format.first_line_indent = Cm(-0.4)
        run = paragraph.add_run(f"{index}) {item}")
        set_run_font(run, size=10.5)


def add_bulleted_items(document: Document, items: list[str]):
    for item in items:
        paragraph = document.add_paragraph()
        set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, after=4)
        paragraph.paragraph_format.left_indent = Cm(0.5)
        paragraph.paragraph_format.first_line_indent = Cm(-0.3)
        run = paragraph.add_run(f"• {item}")
        set_run_font(run, size=10.5)


# ============================================================
# 8. 기관 제출용 Word 보고서 생성
# ============================================================


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = KOREAN_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), KOREAN_FONT)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Cm(17.0))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    set_column_widths(header_table, [9.0, 8.0])
    set_cell_text(
        header_table.cell(0, 0),
        ORGANIZATION_NAME,
        bold=True,
        size=9,
        color=TITLE_COLOR,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    set_cell_text(
        header_table.cell(0, 1),
        "해양 쓰레기 드론 탐지 결과 보고서",
        size=8.5,
        color=DARK_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    set_table_borders(header_table, color=WHITE, size=0)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(17.0))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    set_column_widths(footer_table, [11.5, 5.5])
    set_cell_text(
        footer_table.cell(0, 0),
        f"{ORGANIZATION_NAME} | {DEPARTMENT_NAME}",
        size=8,
        color=DARK_GRAY,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    page_paragraph = footer_table.cell(0, 1).paragraphs[0]
    page_paragraph.text = ""
    add_page_number(page_paragraph)
    set_table_borders(footer_table, color=WHITE, size=0)


def add_document_control_table(document: Document, data: dict[str, Any]):
    table = document.add_table(rows=3, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [2.2, 3.6, 2.2, 3.6, 2.2, 3.6])
    set_table_borders(table, color="9EADBA", size=6)

    values = [
        ("문서번호", display_value(data.get("mission_id")), "작성일", datetime.now().strftime("%Y-%m-%d"), "공개구분", DOCUMENT_SECURITY),
        ("담당부서", DEPARTMENT_NAME, "작성자", AUTHOR_NAME, "조사구역", display_value(data.get("inspection_area"))),
        ("검토", REVIEWER_NAME, "승인", APPROVER_NAME, "조사일", display_value(data.get("inspection_date"))),
    ]

    for row_idx, row_values in enumerate(values):
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            if col_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                set_cell_text(cell, value, bold=True, size=9, color=TITLE_COLOR)
            else:
                set_cell_text(cell, value, size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_title_block(document: Document, data: dict[str, Any]):
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=3, line_spacing=1.0)
    run = paragraph.add_run("해양 쓰레기 드론 탐지 결과 보고서")
    set_run_font(run, size=22, bold=True, color=TITLE_COLOR)

    subtitle = document.add_paragraph()
    set_paragraph_format(subtitle, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=12, line_spacing=1.0)
    subtitle_run = subtitle.add_run(
        f"{display_value(data.get('inspection_area'))} | "
        f"{display_value(data.get('inspection_date'))}"
    )
    set_run_font(subtitle_run, size=11, color=SUBTITLE_COLOR)


def add_executive_summary(document: Document, narrative: dict[str, Any]):
    table = document.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [17.0])
    set_table_borders(table, color=TITLE_COLOR, size=8)

    header_cell = table.cell(0, 0)
    set_cell_shading(header_cell, TITLE_COLOR)
    set_cell_text(
        header_cell,
        "핵심 요약",
        bold=True,
        size=11,
        color=WHITE,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    body_cell = table.cell(1, 0)
    set_cell_shading(body_cell, "F8FAFC")
    set_cell_text(
        body_cell,
        narrative["executive_summary"],
        size=10.5,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    body_cell.paragraphs[0].paragraph_format.line_spacing = 1.45


def add_overview_table(document: Document, data: dict[str, Any], narrative: dict[str, Any] | None = None):
    add_section_heading(document, "1", "조사 개요")
    table = document.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [2.5, 6.0, 2.5, 6.0])
    set_table_borders(table)

    priority = None
    if narrative:
        priority = narrative.get("collection_priority")
    priority = priority or data.get("collection_priority")

    rows = [
        ("미션 ID", display_value(data.get("mission_id")), "비행 ID", display_value(data.get("flight_id"))),
        ("조사 일자", display_value(data.get("inspection_date")), "조사 시간", display_value(data.get("survey_duration_minutes"), "분")),
        ("조사 구역", display_value(data.get("inspection_area")), "조사 방법", display_value(data.get("survey_method"))),
        ("총 탐지 수", display_value(data.get("total_detected_count"), "개"), "수거 우선순위", display_value(priority)),
    ]

    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            cell = table.cell(r_idx, c_idx)
            if c_idx in (0, 2):
                set_cell_shading(cell, LIGHT_GRAY)
                set_cell_text(cell, value, bold=True, color=TITLE_COLOR)
            else:
                set_cell_text(cell, value, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def add_detection_table(document: Document, data: dict[str, Any], narrative: dict[str, Any]):
    add_section_heading(document, "2", "탐지 결과")
    add_body_paragraph(document, narrative["overall_assessment"], indent=True)

    detections = data.get("detections", [])
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [1.1, 3.2, 2.5, 2.0, 2.4, 3.3])
    set_table_borders(table)

    headers = ["번호", "쓰레기 종류", "탐지 구역", "수량", "구성비", "평균 탐지 신뢰도"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, TITLE_COLOR)
        set_cell_text(cell, header, bold=True, color=WHITE, size=9)

    total = int(data.get("total_detected_count") or 0)

    if detections:
        for index, item in enumerate(detections, start=1):
            cells = table.add_row().cells
            values = [
                str(index),
                display_value(item.get("trash_type")),
                display_value(item.get("zone")),
                display_value(item.get("count"), "개"),
                calculate_ratio(int(item.get("count") or 0), total),
                confidence_percent(item.get("average_confidence")),
            ]
            for col_idx, value in enumerate(values):
                alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (1, 2) else WD_ALIGN_PARAGRAPH.CENTER
                set_cell_text(cells[col_idx], value, alignment=alignment)
                if index % 2 == 0:
                    set_cell_shading(cells[col_idx], "F8F8F8")
    else:
        cells = table.add_row().cells
        merged = cells[0]
        for cell in cells[1:]:
            merged = merged.merge(cell)
        set_cell_text(merged, "해당 조사에서 탐지된 해양 쓰레기 객체가 없습니다.", alignment=WD_ALIGN_PARAGRAPH.CENTER)

    note = document.add_paragraph()
    set_paragraph_format(note, alignment=WD_ALIGN_PARAGRAPH.LEFT, before=3, after=2, line_spacing=1.1)
    run = note.add_run("※ 탐지 신뢰도는 객체 탐지 모델의 추정값이며, 현장 확정 판정과 동일하지 않음.")
    set_run_font(run, size=8.5, color=DARK_GRAY)


def add_zone_analysis(document: Document, narrative: dict[str, Any]):
    add_section_heading(document, "3", "구역별 분석")
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [3.2, 13.8])
    set_table_borders(table)

    for idx, header in enumerate(("구역", "분석 내용")):
        set_cell_shading(table.cell(0, idx), LIGHT_BLUE)
        set_cell_text(table.cell(0, idx), header, bold=True, color=TITLE_COLOR)

    for item in narrative["zone_analysis"]:
        cells = table.add_row().cells
        set_cell_text(cells[0], item["zone"], bold=True)
        set_cell_text(cells[1], item["analysis"], alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10)
        cells[1].paragraphs[0].paragraph_format.line_spacing = 1.35


def add_priority_section(document: Document, data: dict[str, Any], narrative: dict[str, Any]):
    add_section_heading(document, "4", "수거 및 현장 확인 우선순위")

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [4.0, 13.0])
    set_table_borders(table, color=TITLE_COLOR, size=8)

    priority_cell = table.cell(0, 0)
    priority = display_value(
        narrative.get("collection_priority") or data.get("collection_priority")
    )
    set_cell_shading(priority_cell, LIGHT_BLUE)
    set_cell_text(priority_cell, f"우선순위\n{priority}", bold=True, size=12, color=TITLE_COLOR)

    set_cell_text(
        table.cell(0, 1),
        narrative["priority_basis"],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        size=10.5,
    )
    table.cell(0, 1).paragraphs[0].paragraph_format.line_spacing = 1.4


def add_actions_section(document: Document, narrative: dict[str, Any]):
    add_section_heading(document, "5", "권고 조치")
    add_numbered_items(document, narrative["recommended_actions"])


def add_limitations_section(document: Document, narrative: dict[str, Any]):
    add_section_heading(document, "6", "데이터 한계 및 유의사항")
    add_bulleted_items(document, narrative["limitations"])


def add_conclusion_section(document: Document, narrative: dict[str, Any]):
    add_section_heading(document, "7", "종합 의견")
    add_body_paragraph(document, narrative["conclusion"], indent=True)


def add_review_section(document: Document):
    add_section_heading(document, "8", "검토 및 확인")
    table = document.add_table(rows=3, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, [5.66, 5.66, 5.66])
    set_table_borders(table)

    headers = ["작성", "검토", "승인"]
    names = [AUTHOR_NAME, REVIEWER_NAME, APPROVER_NAME]

    for idx, header in enumerate(headers):
        set_cell_shading(table.cell(0, idx), LIGHT_GRAY)
        set_cell_text(table.cell(0, idx), header, bold=True, color=TITLE_COLOR)
        set_cell_text(table.cell(1, idx), names[idx], size=10)
        set_cell_text(table.cell(2, idx), "서명 또는 전자결재", size=9, color=DARK_GRAY)
        table.cell(2, idx).height = Cm(1.5)


def build_institutional_docx(
    data: dict[str, Any],
    narrative: dict[str, Any],
    output_path: Path,
    metadata: dict[str, Any] | None = None,
) -> Path:
    document = Document()
    configure_document(document)

    document.core_properties.title = "해양 쓰레기 드론 탐지 결과 보고서"
    document.core_properties.subject = display_value(data.get("inspection_area"))
    document.core_properties.author = AUTHOR_NAME
    document.core_properties.keywords = "해양 쓰레기, 드론, 객체 탐지, 조사 보고서"

    add_document_control_table(document, data)
    add_title_block(document, data)
    add_executive_summary(document, narrative)
    add_overview_table(document, data, narrative)
    add_detection_table(document, data, narrative)
    add_zone_analysis(document, narrative)
    add_priority_section(document, data, narrative)
    add_actions_section(document, narrative)
    add_limitations_section(document, narrative)
    add_conclusion_section(document, narrative)
    add_review_section(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


# ============================================================
# 9. 실행
# ============================================================


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "EXAONE GGUF 모델을 이용해 기관 제출용 "
            "해양 쓰레기 탐지 Word 보고서를 생성합니다."
        )
    )
    parser.add_argument(
        "--case",
        choices=sorted(TEST_CASES.keys()),
        default=None,
        help="내장 테스트 데이터만 사용 (실제 조사 JSON 대신). 지정하지 않으면 logs/surveys 최신 파일 사용",
    )
    parser.add_argument(
        "--input-json",
        type=Path,
        help="특정 조사 JSON 경로 (예: logs/surveys/MISSION-20260721-155113-TT-02.json)",
    )
    parser.add_argument(
        "--latest",
        action="store_true",
        help=f"가장 최근 조사 JSON 사용 (기본 동작과 동일, {SURVEY_DIR})",
    )
    parser.add_argument(
        "--model-path",
        type=Path,
        default=MODEL_PATH,
        help="GGUF 파일 또는 GGUF 파일이 들어 있는 폴더 경로",
    )
    parser.add_argument(
        "--output",
        type=Path,
        help="생성할 DOCX 파일 경로",
    )
    parser.add_argument(
        "--sample-only",
        action="store_true",
        help="모델을 실행하지 않고 기관용 Word 서식만 생성",
    )
    parser.add_argument(
        "--n-gpu-layers",
        type=int,
        default=int(os.getenv("N_GPU_LAYERS", "20")),
        help=(
            "GPU에 올릴 레이어 수. RTX 3060 6GB는 20부터 시작하고, "
            "메모리 부족 시 16/12/8로 낮추세요. -1은 전체 레이어입니다."
        ),
    )
    parser.add_argument(
        "--n-ctx",
        type=int,
        default=int(os.getenv("N_CTX", "4096")),
        help="모델 컨텍스트 길이",
    )
    parser.add_argument(
        "--n-batch",
        type=int,
        default=int(os.getenv("N_BATCH", "128")),
        help="프롬프트 처리 배치 크기. 메모리 부족 시 64로 낮추세요.",
    )
    parser.add_argument(
        "--n-threads",
        type=int,
        default=int(os.getenv("N_THREADS", "0")),
        help="CPU 스레드 수. 0이면 자동 계산",
    )
    parser.add_argument(
        "--chat-format",
        type=str,
        default=os.getenv("CHAT_FORMAT") or None,
        help=(
            "강제로 적용할 llama.cpp 채팅 형식. "
            "기본값은 GGUF 내부 tokenizer.chat_template 자동 사용"
        ),
    )
    parser.add_argument(
        "--verbose-model",
        action="store_true",
        help="llama.cpp 상세 모델 로딩 로그 출력",
    )
    return parser.parse_args()


REQUIRED_SURVEY_KEYS = (
    "mission_id",
    "inspection_date",
    "inspection_area",
    "flight_id",
    "survey_method",
    "total_detected_count",
    "detections",
)


def validate_survey_data(data: dict[str, Any]) -> dict[str, Any]:
    """드론 조사 JSON이 보고서 입력 스키마를 만족하는지 확인·정규화한다."""
    if not isinstance(data, dict):
        raise ValueError("조사 JSON 루트는 객체여야 합니다.")

    missing = [k for k in REQUIRED_SURVEY_KEYS if k not in data]
    if missing:
        raise ValueError(f"조사 JSON에 필수 키가 없습니다: {', '.join(missing)}")

    detections = data.get("detections")
    if not isinstance(detections, list):
        raise ValueError("detections 는 배열이어야 합니다.")

    for i, item in enumerate(detections):
        if not isinstance(item, dict):
            raise ValueError(f"detections[{i}] 는 객체여야 합니다.")
        for key in ("trash_type", "count"):
            if key not in item:
                raise ValueError(f"detections[{i}] 에 '{key}' 가 없습니다.")

    # 선택 키 기본값
    data.setdefault("survey_duration_minutes", None)
    data.setdefault("highest_density_zone", None)
    data.setdefault("collection_priority", None)
    data.setdefault("recommended_actions", [])
    data.setdefault("limitations", [])
    return data


def resolve_latest_survey_json(survey_dir: Path = SURVEY_DIR) -> Path:
    latest = survey_dir / "latest.json"
    if latest.is_file():
        return latest
    if not survey_dir.is_dir():
        raise FileNotFoundError(
            f"조사 JSON 폴더가 없습니다: {survey_dir}\n"
            "드론 스트림을 한 번 켠 뒤 끄면 logs/surveys/ 에 파일이 생성됩니다."
        )
    candidates = sorted(
        survey_dir.glob("MISSION-*.json"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(
            f"조사 JSON이 없습니다: {survey_dir}\n"
            "드론 연결 → 스트림 시작 → 탐지 → 스트림 종료 후 다시 시도하세요."
        )
    return candidates[0]


def load_input_data(args: argparse.Namespace) -> tuple[dict[str, Any], Path | str]:
    """조사 데이터를 로드한다. 반환: (data, source_label)

    우선순위:
      1) --input-json  특정 파일
      2) --case        내장 테스트 데이터
      3) 기본/--latest logs/surveys 최신 MISSION-*.json
    """
    if args.input_json:
        path = args.input_json
        if not path.is_file():
            raise FileNotFoundError(f"조사 JSON 파일을 찾을 수 없습니다: {path}")
        with path.open("r", encoding="utf-8") as file:
            return validate_survey_data(json.load(file)), path

    if args.case:
        return TEST_CASES[args.case], f"TEST_CASES[{args.case}]"

    path = resolve_latest_survey_json()
    with path.open("r", encoding="utf-8") as file:
        return validate_survey_data(json.load(file)), path


def main() -> None:
    args = parse_args()
    try:
        data, source = load_input_data(args)
    except (FileNotFoundError, ValueError, json.JSONDecodeError) as e:
        raise SystemExit(f"[오류] 조사 데이터 로드 실패: {e}") from e

    print(f"입력 데이터: {source}")
    print(f"미션 ID    : {data.get('mission_id')}")
    print(f"탐지 수    : {data.get('total_detected_count')}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    if args.sample_only:
        narrative = deterministic_narrative(data)
        metadata = {
            "elapsed_seconds": 0,
            "input_tokens": 0,
            "generated_tokens": 0,
            "tokens_per_second": 0.0,
            "parse_status": "sample-only",
            "json_mode_used": False,
            "raw_model_output": "",
            "runtime": "template-only",
        }
    else:
        resolved_model_path = resolve_gguf_model_path(args.model_path)

        check_environment(
            resolved_model_path,
            n_gpu_layers=args.n_gpu_layers,
            n_ctx=args.n_ctx,
            n_batch=args.n_batch,
        )

        model = load_model(
            resolved_model_path,
            n_gpu_layers=args.n_gpu_layers,
            n_ctx=args.n_ctx,
            n_batch=args.n_batch,
            n_threads=args.n_threads,
            chat_format=args.chat_format,
            verbose=args.verbose_model,
        )

        narrative, metadata = generate_narrative(model, data)
        metadata["model_path"] = str(resolved_model_path)
        metadata["n_gpu_layers"] = args.n_gpu_layers
        metadata["n_ctx"] = args.n_ctx
        metadata["n_batch"] = args.n_batch

    filename = (
        f"{data.get('mission_id', 'marine_report')}"
        "_institutional_report.docx"
    )
    output_path = args.output or (OUTPUT_DIR / filename)

    build_institutional_docx(
        data=data,
        narrative=narrative,
        output_path=output_path,
        metadata=metadata,
    )

    structured_path = output_path.with_suffix(".structured.json")
    structured_path.write_text(
        json.dumps(
            {
                "input_data": data,
                "narrative": narrative,
                "generation_metadata": metadata,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    raw_output = metadata.get("raw_model_output")
    if raw_output:
        output_path.with_suffix(".raw.txt").write_text(
            str(raw_output),
            encoding="utf-8",
        )

    print("=" * 72)
    print("기관용 Word 보고서 생성 완료")
    print(f"DOCX       : {output_path.resolve()}")
    print(f"구조화 JSON: {structured_path.resolve()}")
    print(f"생성 방식  : {metadata.get('runtime')}")
    print(f"JSON 상태  : {metadata.get('parse_status')}")

    if metadata.get("elapsed_seconds"):
        print(f"생성 시간  : {metadata.get('elapsed_seconds')}초")
        print(f"생성 속도  : {metadata.get('tokens_per_second')} tokens/s")

    print("=" * 72)
    # 웹 API(Reports)가 파싱하기 위한 한 줄 결과
    print(
        "RESULT_JSON:"
        + json.dumps(
            {
                "ok": True,
                "docx": str(output_path.resolve()),
                "structured_json": str(structured_path.resolve()),
                "mission_id": data.get("mission_id"),
                "runtime": metadata.get("runtime"),
                "parse_status": metadata.get("parse_status"),
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
